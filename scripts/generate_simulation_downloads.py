"""Generate learner activity DOCX files from marked sections in a manuscript."""

from copy import deepcopy
from pathlib import Path
from pathlib import PureWindowsPath
import sys
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import etree


ROOT = Path(__file__).resolve().parents[1]
IMAGE_DIR = ROOT / "assets" / "images"
WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": WORD_NS}
FOOTER_PREFIX = "From J.M. Larson and H.L. Stedge, "
FOOTER_TITLE = "Clinical Simulations for the Athletic Trainer HKPropel Access"
FOOTER_SUFFIX = " (Human Kinetics, 2027)."
FOOTER_TEXT = FOOTER_PREFIX + FOOTER_TITLE + FOOTER_SUFFIX
FOOTER_SIZE = Pt(9)
REQUIRED_FONT = "Aptos"
REQUIRED_HEADING_SIZE = Pt(20)
SIMULATIONS = (
    {
        "id": "simulation-1",
        "source": ROOT / "assets" / "Manuscripts" / "E9814_Sim01_Time-Out Ankle.docx",
        "page": ROOT / "pages" / "simulation-1.html",
        "output_dir": ROOT / "assets" / "downloads" / "simulation-1",
        "activities": (
            ("Presimulation Activity 2: Anatomy Labeling", "Presimulation Activity 2: Anatomy Labeling", "presimulation-activity-2-anatomy-labeling.docx"),
            ("Presimulation Activity 3: Wrestling Information", "Presimulation Activity 3: Wrestling Information", "presimulation-activity-3-wrestling-information.docx"),
            ("Presimulation Activity 4: Evaluation Efficiency", "Presimulation Activity 4: Evaluation Efficiency", "presimulation-activity-4-evaluation-efficiency.docx"),
            ("Presimulation Activity 5: Ankle Taping", "Presimulation Activity 5: Ankle Taping", "presimulation-activity-5-ankle-taping.docx"),
            ("Presimulation Activity 6: Discussing RTP", "Presimulation Activity 6: Discussing RTP", "presimulation-activity-6-discussing-rtp.docx"),
            ("Presimulation Activity 7: Sideline Kit Packing", "Presimulation Activity 7: Sideline Kit Packing", "presimulation-activity-7-sideline-kit-packing.docx"),
            ("Postsimulation Activity 1: Documentation", "Postsimulation Activity 1: Documentation", "postsimulation-activity-1-documentation.docx"),
            ("Postsimulation Activity 2: Reflection", "Postsimulation Activity 2: Reflection", "postsimulation-activity-2-reflection.docx"),
        ),
    },
    {
        "id": "simulation-2",
        "source": ROOT / "assets" / "Manuscripts" / "E9814_Sim02 Exercise Illness.docx",
        "page": ROOT / "pages" / "simulation-2.html",
        "output_dir": ROOT / "assets" / "downloads" / "simulation-2",
        "activities": (
            ("Activity 2: List", "Presimulation Activity 2: List", "presimulation-activity-2-list.docx"),
            ("Activity 3: Body System", "Presimulation Activity 3: Body System", "presimulation-activity-3-body-system.docx"),
            ("Activity 4: History", "Presimulation Activity 4: History", "presimulation-activity-4-history.docx"),
            ("Postsimulation Activity 1: Review", "Postsimulation Activity 1: Review", "postsimulation-activity-1-review.docx"),
        ),
    },
    {
        "id": "simulation-3",
        "source": ROOT / "assets" / "Manuscripts" / "E9814_Sim03_EAP.docx",
        "page": ROOT / "pages" / "simulation-3.html",
        "output_dir": ROOT / "assets" / "downloads" / "simulation-3",
        "activities": (
            ("Presimulation Activity 2: EAP Review", "Presimulation Activity 2: EAP Review", "presimulation-activity-2-eap-review.docx"),
            ("Presimulation Activity 3: Roles and Responsibilities", "Presimulation Activity 3: Roles and Responsibilities ", "presimulation-activity-3-roles-and-responsibilities.docx"),
            ("Presimulation Activity 4: Practice", "Presimulation Activity 4: Practice", "presimulation-activity-4-practice.docx"),
            ("Presimulation Activity 5: Communication", "Presimulation Activity 5: Communication", "presimulation-activity-5-communication.docx"),
            ("Presimulation Activity 6: Transfer or Care", "Presimulation Activity 6: Transfer of Care", "presimulation-activity-6-transfer-of-care.docx"),
        ),
    },
    {
        "id": "simulation-4",
        "source": ROOT / "assets" / "Manuscripts" / "E9814_Sim04_Special Tests-LE.docx",
        "page": ROOT / "pages" / "simulation-4.html",
        "output_dir": ROOT / "assets" / "downloads" / "simulation-4",
        "activities": (
            ("Presimulation Activity 2: Palpation", "Presimulation Activity 2: Palpation", "presimulation-activity-2-palpation.docx"),
            ("Presimulation Activity 3: Anatomy", "Presimulation Activity 3: Anatomy", "presimulation-activity-3-anatomy.docx"),
            ("Presimulation Activity 4: Special Test Understandings", "Presimulation Activity 4: Special Test Understandings", "presimulation-activity-4-special-test-understandings.docx"),
            ("Presimulation Activity 5 (Required): Special Test Selection", "Presimulation Activity 5: Special Test Selection", "presimulation-activity-5-special-test-selection.docx"),
            ("Postsimulation Activity 1: Special Test Rationale", "Postsimulation Activity 1: Special Test Rationale", "postsimulation-activity-1-special-test-rationale.docx"),
            ("Postsimulation Activity 2: Mechanism of Injuries", "Postsimulation Activity 2: Mechanism of Injury", "postsimulation-activity-2-mechanism-of-injury.docx"),
            ("Postsimulation Activity 3: Plan of Care", "Postsimulation Activity 3: Plan of Care", "postsimulation-activity-3-plan-of-care.docx"),
            ("Postsimulation Activity 4: Diagnostic Accuracy", "Postsimulation Activity 4: Diagnostic Accuracy ", "postsimulation-activity-4-diagnostic-accuracy.docx"),
        ),
    },
    {
        "id": "simulation-5",
        "source": ROOT / "assets" / "Manuscripts" / "E9814_Sim05_Special Tests_UE.docx",
        "page": ROOT / "pages" / "simulation-5.html",
        "output_dir": ROOT / "assets" / "downloads" / "simulation-5",
        "activities": (
            ("Presimulation Activity 2: Palpation", "Presimulation Activity 2: Palpation", "presimulation-activity-2-palpation.docx"),
            ("Presimulation Activity 3: Background Information", "Presimulation Activity 3: Background Information", "presimulation-activity-3-background-information.docx"),
            ("Presimulation Activity 4: Anatomy", "Presimulation Activity 4: Anatomy", "presimulation-activity-4-anatomy.docx"),
            ("Presimulation Activity 5: Special Test Understandings", "Presimulation Activity 5: Special Test Understandings", "presimulation-activity-5-special-test-understandings.docx"),
            ("Activity 6 (Required): Special Test Selection", "Presimulation Activity 6: Special Test Selection", "presimulation-activity-6-special-test-selection.docx"),
            ("Postsimulation Activity 1: Special Test Rationale", "Postsimulation Activity 1: Special Test Rationale", "postsimulation-activity-1-special-test-rationale.docx"),
            ("Postsimulation Activity 2: Mechanisms of Injuries", "Postsimulation Activity 2: Mechanisms of Injuries", "postsimulation-activity-2-mechanisms-of-injuries.docx"),
            ("Postsimulation Activity 3: Plan of Care", "Postsimulation Activity 3: Plan of Care", "postsimulation-activity-3-plan-of-care.docx"),
        ),
    },
    {
        "id": "simulation-6",
        "source": ROOT / "assets" / "Manuscripts" / "E9814_Sim06_Dyspnea.docx",
        "page": ROOT / "pages" / "simulation-6.html",
        "output_dir": ROOT / "assets" / "downloads" / "simulation-6",
        "activities": (
            ("Presimulation Activity 2: Signs and Symptoms", "Presimulation Activity 2: Signs and Symptoms", "presimulation-activity-2-signs-and-symptoms.docx"),
            ("Presimulation Activity 3: Pertinent Negative Patterns", "Presimulation Activity 3: Pertinent Negative Patterns", "presimulation-activity-3-pertinent-negative-patterns.docx"),
            ("Presimulation Activity 4: Action Plan", "Presimulation Activity 4: Action Plan", "presimulation-activity-4-action-plan.docx"),
            ("Postsimulation Activity 1: Asthma Action Plan", "Postsimulation Activity 1: Asthma Action Plan", "postsimulation-activity-1-asthma-action-plan.docx"),
            ("Postsimulation Activity 2: Anaphylaxis Action Plan", "Postsimulation Activity 2: Anaphylaxis Action Plan", "postsimulation-activity-2-anaphylaxis-action-plan.docx"),
            ("Postsimulation Activity 3: Reflection", "Postsimulation Activity 3: Reflection", "postsimulation-activity-3-reflection.docx"),
        ),
    },
    {
        "id": "simulation-7",
        "source": ROOT / "assets" / "Manuscripts" / "E9814_Sim07_Phases of Healing.docx",
        "page": ROOT / "pages" / "simulation-7.html",
        "output_dir": ROOT / "assets" / "downloads" / "simulation-7",
        "activities": (
            ("Presimulation Activity 2: Describing the Phases of Healing", "Presimulation Activity 2: Describing the Phases of Healing", "presimulation-activity-2-describing-the-phases-of-healing.docx"),
            ("Presimulation Activity 3: Communication", "Presimulation Activity 3: Communication", "presimulation-activity-3-communication.docx"),
            ("Presimulation Activity 4: Anatomy", "Presimulation Activity 4: Anatomy", "presimulation-activity-4-anatomy.docx"),
            ("Postsimulation Activity 1: Ethics and Values", "Postsimulation Activity 1: Ethics and Values", "postsimulation-activity-1-ethics-and-values.docx"),
            ("Postsimulation Activity 2: Reflection", "Postsimulation Activity 2: Reflection", "postsimulation-activity-2-reflection.docx"),
            ("Postsimulation Activity 3: Boundaries Interview", "Postsimulation Activity 3: Boundaries Interview", "postsimulation-activity-3-boundaries-interview.docx"),
        ),
    },
    {
        "id": "simulation-8",
        "source": ROOT / "assets" / "Manuscripts" / "E9814_Sim08_Abdominal Injury.docx",
        "page": ROOT / "pages" / "simulation-8.html",
        "output_dir": ROOT / "assets" / "downloads" / "simulation-8",
        "activities": (
            ("Presimulation Activity 2: Clinical Observations and Assessments", "Presimulation Activity 2: Clinical Observations and Assessments", "presimulation-activity-2-clinical-observations-and-assessments.docx"),
            ("Presimulation Activity 3: Anatomy Labeling", "Presimulation Activity 3: Anatomy Labeling", "presimulation-activity-3-anatomy-labeling.docx"),
            ("Postsimulation Activity 1: Pertinent Negatives", "Postsimulation Activity 1: Pertinent Negatives", "postsimulation-activity-1-pertinent-negatives.docx"),
            ("Postsimulation Activity 2: Reflection", "Postsimulation Activity 2: Reflection", "postsimulation-activity-2-reflection.docx"),
        ),
    },
    {
        "id": "simulation-9",
        "source": ROOT / "assets" / "Manuscripts" / "E9814_Sim09_Coach Education.docx",
        "page": ROOT / "pages" / "simulation-9.html",
        "output_dir": ROOT / "assets" / "downloads" / "simulation-9",
        "activities": (
            ("Presimulation Activity 2: Condition Template", "Presimulation Activity 2: Condition Template", "presimulation-activity-2-condition-template.docx"),
            ("Presimulation Activity 3: Communication Strategies", "Presimulation Activity 3: Communication Strategies", "presimulation-activity-3-communication-strategies.docx"),
            ("Presimulation Activity 4: Emergency Action Plan Review", "Presimulation Activity 4: Emergency Action Plan Review", "presimulation-activity-4-emergency-action-plan-review.docx"),
            ("Postsimulation Activity 1: Summarizing Email", "Postsimulation Activity 1: Summarizing Email", "postsimulation-activity-1-summarizing-email.docx"),
        ),
    },
    {
        "id": "simulation-10",
        "source": ROOT / "assets" / "Manuscripts" / "E9814_Sim10_History Assessment.docx",
        "page": ROOT / "pages" / "simulation-10.html",
        "output_dir": ROOT / "assets" / "downloads" / "simulation-10",
        "activities": (
            ("Presimulation Activity 2: SAMPLE OPQRST", "Presimulation Activity 2: SAMPLE OPQRST", "presimulation-activity-2-sample-opqrst.docx"),
            ("Presimulation Activity 3: Condition History With a Friend", "Presimulation Activity 3: Condition History With a Friend", "presimulation-activity-3-condition-history-with-a-friend.docx"),
            ("Presimulation Activity 4: Take Two Histories", "Presimulation Activity 4: Take Two Histories", "presimulation-activity-4-take-two-histories.docx"),
            ("Presimulation Activity 5: Telephone Game", "Presimulation Activity 5: Telephone Game", "presimulation-activity-5-telephone-game.docx"),
            ("Postsimulation Activity 1: Structured History", "Postsimulation Activity 1: Structured History", "postsimulation-activity-1-structured-history.docx"),
            ("Postsimulation Activity 2: Reflection", "Postsimulation Activity 2: Reflection", "postsimulation-activity-2-reflection.docx"),
        ),
    },
    {
        "id": "simulation-11",
        "source": ROOT / "assets" / "Manuscripts" / "E9814_Sim11_Modalities Shuffle Cold and Hot.docx",
        "page": ROOT / "pages" / "simulation-11.html",
        "output_dir": ROOT / "assets" / "downloads" / "simulation-11",
        "activities": (
            ("Presimulation Activity 2: Contraindications and Precautions", "Presimulation Activity 2: Contraindications and Precautions", "presimulation-activity-2-contraindications-and-precautions.docx"),
            ("Presimulation Activity 3: Indications", "Presimulation Activity 3: Indications", "presimulation-activity-3-indications.docx"),
            ("Presimulation Activity 4: Physiological Effects", "Presimulation Activity 4: Physiological Effects", "presimulation-activity-4-physiological-effects.docx"),
            ("Presimulation Activity 5: Application", "Presimulation Activity 5: Application", "presimulation-activity-5-application.docx"),
            ("Postsimulation Activity 1: Electronic Medical Record", "Postsimulation Activity 1: Electronic Medical Record and Reflection", "postsimulation-activity-1-electronic-medical-record-and-reflection.docx"),
            ("Postsimulation Activity 2: SP", "Postsimulation Activity 2: SP", "postsimulation-activity-2-sp.docx"),
            ("Postsimulation Activity 3: Proctor", "Postsimulation Activity 3: Proctor", "postsimulation-activity-3-proctor.docx"),
        ),
    },
    {
        "id": "simulation-12",
        "source": ROOT / "assets" / "Manuscripts" / "E9814_Sim12_Modalities Shuffle Electro and Ultrasound.docx",
        "page": ROOT / "pages" / "simulation-12.html",
        "output_dir": ROOT / "assets" / "downloads" / "simulation-12",
        "activities": (
            ("Presimulation Activity 2: Indications and Contraindications", "Presimulation Activity 2: Indications and Contraindications", "presimulation-activity-2-indications-and-contraindications.docx"),
            ("Presimulation Activity 3: Physiological Effects", "Presimulation Activity 3: Physiological Effects", "presimulation-activity-3-physiological-effects.docx"),
            ("Presimulation Activity 4: Practice", "Presimulation Activity 4: Practice", "presimulation-activity-4-practice.docx"),
            ("Postsimulation Activity 1: Selection Rationale", "Postsimulation Activity 1: Selection Rationale", "postsimulation-activity-1-selection-rationale.docx"),
            ("Postsimulation Activity 2: Documentation", "Postsimulation Activity 2: Documentation", "postsimulation-activity-2-documentation.docx"),
        ),
    },
    {
        "id": "simulation-13",
        "source": ROOT / "assets" / "Manuscripts" / "E9814_Sim13_Shoulder Injury.docx",
        "page": ROOT / "pages" / "simulation-13.html",
        "output_dir": ROOT / "assets" / "downloads" / "simulation-13",
        "activities": (
            ("Presimulation Activity 2: Annotation", "Presimulation Activity 2: Annotation", "presimulation-activity-2-annotation.docx"),
            ("Presimulation Activity 3: Palpation and Anatomy", "Presimulation Activity 3: Palpation and Anatomy", "presimulation-activity-3-palpation-and-anatomy.docx"),
            ("Presimulation Activity 4: Readiness Questions", "Presimulation Activity 4: Readiness Questions", "presimulation-activity-4-readiness-questions.docx"),
            ("Presimulation Activity 5: Partner Practice", "Presimulation Activity 5: Partner Practice", "presimulation-activity-5-partner-practice.docx"),
            ("Postsimulation Activity 1: Reflection", "Postsimulation Activity 1: Reflection", "postsimulation-activity-1-reflection.docx"),
        ),
    },
    {
        "id": "simulation-14",
        "source": ROOT / "assets" / "Manuscripts" / "E9814_Sim14_Taping Ethics.docx",
        "page": ROOT / "pages" / "simulation-14.html",
        "output_dir": ROOT / "assets" / "downloads" / "simulation-14",
        "activities": (
            ("Presimulation Activity 2: Ethical Scenario 1", "Presimulation Activity 2: Ethical Scenario 1", "presimulation-activity-2-ethical-scenario-1.docx"),
            ("Presimulation Activity 3: Ethical Scenario 2", "Presimulation Activity 3: Ethical Scenario 2", "presimulation-activity-3-ethical-scenario-2.docx"),
            ("Presimulation Activity 4: Ethical Scenario 3", "Presimulation Activity 4: Ethical Scenario 3", "presimulation-activity-4-ethical-scenario-3.docx"),
            ("Presimulation Activity 5: Ethical Scenario 4", "Presimulation Activity 5: Ethical Scenario 4", "presimulation-activity-5-ethical-scenario-4.docx"),
            ("Presimulation Activity 6: Ethical Scenario 5", "Presimulation Activity 6: Ethical Scenario 5", "presimulation-activity-6-ethical-scenario-5.docx"),
            ("Presimulation Activity 7: Ethical Scenario 6", "Presimulation Activity 7: Ethical Scenario 6", "presimulation-activity-7-ethical-scenario-6.docx"),
            ("Presimulation Activity 8: Ethical Scenario 7", "Presimulation Activity 8: Ethical Scenario 7", "presimulation-activity-8-ethical-scenario-7.docx"),
            ("Presimulation Activity 9: Ethical Scenario 8", "Presimulation Activity 9: Ethical Scenario 8", "presimulation-activity-9-ethical-scenario-8.docx"),
            ("Presimulation Activity 10: Ethical Scenario 9", "Presimulation Activity 10: Ethical Scenario 9", "presimulation-activity-10-ethical-scenario-9.docx"),
            ("Presimulation Activity 11: Ethical Scenario 10", "Presimulation Activity 11: Ethical Scenario 10", "presimulation-activity-11-ethical-scenario-10.docx"),
            ("Presimulation Activity 12: Resources", "Presimulation Activity 12: Resources", "presimulation-activity-12-resources.docx"),
            ("Postsimulation Activity 1: Reflection", "Postsimulation Activity 1: Reflection", "postsimulation-activity-1-reflection.docx"),
        ),
    },
    {
        "id": "simulation-15",
        "source": ROOT / "assets" / "Manuscripts" / "E9814_Sim15_Primary Assessment.docx",
        "page": ROOT / "pages" / "simulation-15.html",
        "output_dir": ROOT / "assets" / "downloads" / "simulation-15",
        "activities": (
            ("Presimulation Activity 2: List", "Presimulation Activity 2: List", "presimulation-activity-2-list.docx"),
            ("Presimulation Activity 3: Partner List (Presimulation Activity 2 Required First)", "Presimulation Activity 3: Partner List (Presimulation Activity 2 Required First)", "presimulation-activity-3-partner-list.docx"),
            ("Presimulation Activity 4: Primary Assessment Practice", "Presimulation Activity 4: Primary Assessment Practice", "presimulation-activity-4-primary-assessment-practice.docx"),
            ("Presimulation Activity 5: Identify Components", "Presimulation Activity 5: Identify Components", "presimulation-activity-5-identify-components.docx"),
            ("Postsimulation Activity 1: Pulse and Heart Rate", "Postsimulation Activity 1: Pulse and Heart Rate", "postsimulation-activity-1-pulse-and-heart-rate.docx"),
            ("Postsimulation Activity 2: Oxygen and Respiration", "Postsimulation Activity 2: Oxygen and Respiration", "postsimulation-activity-2-oxygen-and-respiration.docx"),
            ("Postsimulation Activity 3: Blood Pressure and Perfusion", "Postsimulation Activity 3: Blood Pressure and Perfusion", "postsimulation-activity-3-blood-pressure-and-perfusion.docx"),
            ("Presimulation Activity 4: Other Assessments", "Presimulation Activity 4: Other Assessments", "presimulation-activity-4-other-assessments.docx"),
        ),
    },
    {
        "id": "simulation-16",
        "source": ROOT / "assets" / "Manuscripts" / "E9814_Sim16_ A Good Mimic.docx",
        "page": ROOT / "pages" / "simulation-16.html",
        "output_dir": ROOT / "assets" / "downloads" / "simulation-16",
        "activities": (
            ("Presimulation Activity 2: Medications", "Presimulation Activity 2: Medications", "presimulation-activity-2-medications.docx"),
            ("Presimulation Activity 3: Injury Evaluation", "Presimulation Activity 3: Injury Evaluation", "presimulation-activity-3-injury-evaluation.docx"),
            ("Presimulation Activity 4: Power Wheel", "Presimulation Activity 4: Power Wheel", "presimulation-activity-4-power-wheel.docx"),
            ("Postsimulation Activity 1: Speaking Out", "Postsimulation Activity 1: Speaking Out", "postsimulation-activity-1-speaking-out.docx"),
            ("Postsimulation Activity 2: Medication Precautions", "Postsimulation Activity 2: Medication Precautions", "postsimulation-activity-2-medication-precautions.docx"),
            ("Postsimulation Activity 3: Pertinent Negatives and Positives", "Postsimulation Activity 3: Pertinent Negatives and Positives", "postsimulation-activity-3-pertinent-negatives-and-positives.docx"),
            ("Postsimulation Activity 4: Commentary", "Postsimulation Activity 4: Commentary", "postsimulation-activity-4-commentary.docx"),
        ),
    },
    {
        "id": "simulation-17",
        "source": ROOT / "assets" / "Manuscripts" / "E9814_Sim17_SSIMS.docx",
        "page": ROOT / "pages" / "simulation-17.html",
        "output_dir": ROOT / "assets" / "downloads" / "simulation-17",
        "activities": (
            ("Presimulation Activity 2: Special Tests", "Presimulation Activity 2: Special Tests", "presimulation-activity-2-special-tests.docx"),
            ("Presimulation Activity 3: Anatomy", "Presimulation Activity 3: Anatomy", "presimulation-activity-3-anatomy.docx"),
            ("Presimulation Activity 4: Insurance", "Presimulation Activity 4: Insurance", "presimulation-activity-4-insurance.docx"),
            ("Presimulation Activity 5: Social Determinants of Health", "Presimulation Activity 5: Social Determinants of Health", "presimulation-activity-5-social-determinants-of-health.docx"),
            ("Postsimulation Activity 1: Reflection", "Postsimulation Activity 1: Reflection", "postsimulation-activity-1-reflection.docx"),
            ("Postsimulation Activity 2: Feedback on Media", "Postsimulation Activity 2: Feedback on Media", "postsimulation-activity-2-feedback-on-media.docx"),
            ("Postsimulation Activity 3: Self-Assessment", "Postsimulation Activity 3: Self-Assessment", "postsimulation-activity-3-self-assessment.docx"),
            ("Postsimulation Activity 4: Change", "Postsimulation Activity 4: Change", "postsimulation-activity-4-change.docx"),
        ),
    },
    {
        "id": "simulation-18",
        "source": ROOT / "assets" / "Manuscripts" / "E9814_Sim18_CSI.docx",
        "page": ROOT / "pages" / "simulation-18.html",
        "output_dir": ROOT / "assets" / "downloads" / "simulation-18",
        "activities": (
            ("Presimulation Activity 2: Emergency Action Plan (EAP)", "Presimulation Activity 2: Emergency Action Plan (EAP)", "presimulation-activity-2-emergency-action-plan.docx"),
            ("Presimulation Activity 3: Practice", "Presimulation Activity 3: Practice", "presimulation-activity-3-practice.docx"),
            ("Presimulation Activity 4: Communication", "Presimulation Activity 4: Communication", "presimulation-activity-4-communication.docx"),
            ("Postsimulation Activity 1: Elevator Speech", "Postsimulation Activity 1: Elevator Speech", "postsimulation-activity-1-elevator-speech.docx"),
            ("Activity 2: Roles and Actions", "Postsimulation Activity 2: Roles and Actions", "postsimulation-activity-2-roles-and-actions.docx"),
        ),
    },
    {
        "id": "simulation-19",
        "source": ROOT / "assets" / "Manuscripts" / "E9814_Sim19_Snowboard Accident.docx",
        "page": ROOT / "pages" / "simulation-19.html",
        "output_dir": ROOT / "assets" / "downloads" / "simulation-19",
        "activities": (
            ("Presimulation Activity 2: Emergency Action Plan (EAP)", "Presimulation Activity 2: Emergency Action Plan (EAP)", "presimulation-activity-2-emergency-action-plan.docx"),
            ("Presimulation Activity 3: Primary Assessment", "Presimulation Activity 3: Primary Assessment", "presimulation-activity-3-primary-assessment.docx"),
            ("Presimulation Activity 4: Communication", "Presimulation Activity 4: Communication", "presimulation-activity-4-communication.docx"),
            ("Presimulation Activity 5: Equipment Considerations", "Presimulation Activity 5: Equipment Considerations", "presimulation-activity-5-equipment-considerations.docx"),
            ("Presimulation Activity 6: Identification", "Presimulation Activity 6: Identification", "presimulation-activity-6-identification.docx"),
            ("Presimulation Activity 7: Care Plan", "Presimulation Activity 7: Care Plan", "presimulation-activity-7-care-plan.docx"),
            ("Postsimulation Activity 1: After Action Report", "Postsimulation Activity 1: After Action Report", "postsimulation-activity-1-after-action-report.docx"),
            ("Postsimulation Activity 2: Subjective, Objective, Assessment, Plan (SOAP) Note", "Postsimulation Activity 2: Subjective, Objective, Assessment, Plan (SOAP) Note", "postsimulation-activity-2-soap-note.docx"),
            ("Postsimulation Activity 3: Reflection", "Postsimulation Activity 3: Reflection", "postsimulation-activity-3-reflection.docx"),
        ),
    },
    {
        "id": "simulation-20",
        "source": ROOT / "assets" / "Manuscripts" / "E9814_Sim20 Head-to-Head (Head Injury Part 1).docx",
        "page": ROOT / "pages" / "simulation-20.html",
        "output_dir": ROOT / "assets" / "downloads" / "simulation-20",
        "activities": (
            ("Presimulation Activity 2: Head Injury Assessment", "Presimulation Activity 2: Head Injury Assessment", "presimulation-activity-2-head-injury-assessment.docx"),
            ("Presimulation Activity 3: Head Injuries", "Presimulation Activity 3: Head Injuries", "presimulation-activity-3-head-injuries.docx"),
            ("Presimulation Activity 4: Cranial Nerve Routine", "Presimulation Activity 4: Cranial Nerve Routine", "presimulation-activity-4-cranial-nerve-routine.docx"),
            ("Postsimulation Activity 1: Documentation", "Postsimulation Activity 1: Documentation", "postsimulation-activity-1-documentation.docx"),
            ("Postsimulation Activity 2: Preparation", "Postsimulation Activity 2: Preparation", "postsimulation-activity-2-preparation.docx"),
            ("Postsimulation Activity 3: Review of Literature", "Postsimulation Activity 3: Review of Literature", "postsimulation-activity-3-review-of-literature.docx"),
            ("Postsimulation Activity 4: Quiz", "Postsimulation Activity 4: Quiz", "postsimulation-activity-4-quiz.docx"),
        ),
    },
    {
        "id": "simulation-21",
        "source": ROOT / "assets" / "Manuscripts" / "E9814_Sim21 SCAT (Head Injury Part 2).docx",
        "page": ROOT / "pages" / "simulation-21.html",
        "output_dir": ROOT / "assets" / "downloads" / "simulation-21",
        "resources": ("SCAT6 Instructions", "HCP SCAT6 Rain Shoemaker"),
        "activities": (
            ("Presimulation Activity 2: Administer SCAT6", "Presimulation Activity 2: Administer SCAT6", "presimulation-activity-2-administer-scat6.docx"),
            ("Presimulation Activity 3: Take-Home Instructions", "Presimulation Activity 3: Take-Home Instructions", "presimulation-activity-3-take-home-instructions.docx"),
            ("Presimulation Activity 4: SCAT6 Instructions", "Presimulation Activity 4: SCAT6 Instructions", "presimulation-activity-4-scat6-instructions.docx"),
            ("Presimulation Activity 5: RTP and RTL", "Presimulation Activity 5: RTP and RTL", "presimulation-activity-5-rtp-and-rtl.docx"),
            ("Rubric", "Rubric: Sport Concussion Assessment Tool (Concussion Part 2)", "rubric-sport-concussion-assessment-tool.docx"),
            ("Postsimulation Activity 1: Patient Care Report (PCR)", "Postsimulation Activity 1: Patient Care Report (PCR)", "postsimulation-activity-1-patient-care-report.docx"),
            ("Postsimulation Activity 2: Plan", "Postsimulation Activity 2: Plan", "postsimulation-activity-2-plan.docx"),
            ("Postsimulation Activity 3: Research and Understanding", "Postsimulation Activity 3: Research and Understanding", "postsimulation-activity-3-research-and-understanding.docx"),
        ),
    },
    {
        "id": "simulation-22",
        "source": ROOT / "assets" / "Manuscripts" / "E9814_Sim22 Provider Communication (Head Injury Part 3).docx",
        "page": ROOT / "pages" / "simulation-22.html",
        "output_dir": ROOT / "assets" / "downloads" / "simulation-22",
        "resources": ("96-Hour SCAT6",),
        "activities": (
            ("Presimulation Activity 2: Referral", "Presimulation Activity 2: Referral", "presimulation-activity-2-referral.docx"),
            ("Presimulation Activity 3: Documentation", "Presimulation Activity 3: Documentation", "presimulation-activity-3-documentation.docx"),
            ("Presimulation Activity 4: Why Refer", "Presimulation Activity 4: Why Refer", "presimulation-activity-4-why-refer.docx"),
            ("Presimulation Activity 5: Why a PT", "Presimulation Activity 5: Why a PT", "presimulation-activity-5-why-a-pt.docx"),
            ("Postsimulation Activity 1: Documentation", "Postsimulation Activity 1: Documentation", "postsimulation-activity-1-documentation.docx"),
            ("Postsimulation Activity 2: Summarizing and Clarifying Email", "Postsimulation Activity 2: Summarizing and Clarifying Email", "postsimulation-activity-2-summarizing-and-clarifying-email.docx"),
            ("Postsimulation Activity 3: Literature Review", "Postsimulation Activity 3: Literature Review", "postsimulation-activity-3-literature-review.docx"),
            ("Physical Therapy and AT/HCP Diagnosis and Concussion Management", "Postsimulation Activity 4: Physical Therapy and AT/HCP Diagnosis and Concussion Management", "postsimulation-activity-4-physical-therapy-and-diagnosis.docx"),
            ("Postsimulation Activity 5: New Knowledge", "Postsimulation Activity 5: New Knowledge", "postsimulation-activity-5-new-knowledge.docx"),
        ),
    },
    {
        "id": "simulation-23", "source": ROOT / "assets" / "Manuscripts" / "E9814_Sim23 EHS.docx", "page": ROOT / "pages" / "simulation-23.html", "output_dir": ROOT / "assets" / "downloads" / "simulation-23",
        "activities": (
            ("Presimulation Activity 2: Reflection Questions", "Presimulation Activity 2: Reflection Questions", "presimulation-activity-2-reflection-questions.docx"),
            ("Presimulation Activity 3: Practice", "Presimulation Activity 3: Practice", "presimulation-activity-3-practice.docx"),
            ("Postsimulation Activity 1: Reflection Paper", "Postsimulation Activity 1: Reflection Paper", "postsimulation-activity-1-reflection-paper.docx"),
            ("Postsimulation Activity 2: Documentation", "Postsimulation Activity 2: Documentation", "postsimulation-activity-2-documentation.docx"),
        ),
    },
    {
        "id": "simulation-24", "source": ROOT / "assets" / "Manuscripts" / "E9814_Sim24 Rotator Cuff Complication.docx", "page": ROOT / "pages" / "simulation-24.html", "output_dir": ROOT / "assets" / "downloads" / "simulation-24",
        "activities": (
            ("Presimulation Activity 2: Signs and Symptoms", "Presimulation Activity 2: Signs and Symptoms", "presimulation-activity-2-signs-and-symptoms.docx"),
            ("Presmulation Activity 3: Surgical Factors", "Presmulation Activity 3: Surgical Factors", "presmulation-activity-3-surgical-factors.docx"),
            ("Presimulation Activity 4: History", "Presimulation Activity 4: History", "presimulation-activity-4-history.docx"),
            ("Presimulation Activity 5: Postsurgical Complications", "Presimulation Activity 5: Postsurgical Complications", "presimulation-activity-5-postsurgical-complications.docx"),
            ("Postsimulation Activity 1: Reflection Paper", "Postsimulation Activity 1: Reflection Paper", "postsimulation-activity-1-reflection-paper.docx"),
            ("Postsimulation Activity 2: Questions Remaining", "Postsimulation Activity 2: Questions Remaining", "postsimulation-activity-2-questions-remaining.docx"),
        ),
    },
    {
        "id": "simulation-25", "source": ROOT / "assets" / "Manuscripts" / "E9814_Sim25 Foot Pain Telehealth.docx", "page": ROOT / "pages" / "simulation-25.html", "output_dir": ROOT / "assets" / "downloads" / "simulation-25",
        "activities": (
            ("Presimulation Activity 2: Understanding Foot Injuries", "Presimulation Activity 2: Understanding Foot Injuries", "presimulation-activity-2-understanding-foot-injuries.docx"),
            ("Presimulation Activity 3: Cardiopulmonary Resuscitation (CPR)", "Presimulation Activity 3: Cardiopulmonary Resuscitation (CPR)", "presimulation-activity-3-cardiopulmonary-resuscitation.docx"),
            ("Presimulation Activity 4: Phone Etiquette", "Presimulation Activity 4: Phone Etiquette", "presimulation-activity-4-phone-etiquette.docx"),
            ("Postsimulation Activity 1: Parent Summary of Athlete Interaction", "Postsimulation Activity 1: Parent Summary of Athlete Interaction", "postsimulation-activity-1-parent-summary.docx"),
            ("Presimulation Activity 2: Reflection", "Presimulation Activity 2: Reflection", "presimulation-activity-2-reflection.docx"),
        ),
    },
    {
        "id": "simulation-26", "source": ROOT / "assets" / "Manuscripts" / "E9814_Sim26 Foot Pain Gait Analysis.docx", "page": ROOT / "pages" / "simulation-26.html", "output_dir": ROOT / "assets" / "downloads" / "simulation-26", "resources": ("Running Gait Videos",),
        "activities": (
            ("Presimulation Activity 2: Phases of Gait", "Presimulation Activity 2: Phases of Gait", "presimulation-activity-2-phases-of-gait.docx"),
            ("Presimulation Activity 3: Functional Analysis", "Presimulation Activity 3: Functional Analysis", "presimulation-activity-3-functional-analysis.docx"),
            ("Presimulation Activity 4: Peer Recorded Gait Analysis", "Presimulation Activity 4: Peer Recorded Gait Analysis", "presimulation-activity-4-peer-recorded-gait-analysis.docx"),
            ("Presimulation Activity 5: Standardized Patient Gait Analysis", "Presimulation Activity 5: Standardized Patient Gait Analysis", "presimulation-activity-5-standardized-patient-gait-analysis.docx"),
            ("Presimulation Activity 6: Video Analysis", "Presimulation Activity 6: Video Analysis", "presimulation-activity-6-video-analysis.docx"),
            ("Postsimulation Activity 1: Analysis and Modifications", "Postsimulation Activity 1: Analysis and Modifications", "postsimulation-activity-1-analysis-and-modifications.docx"),
            ("Postsimulation Activity 2: Home Exercise Program (HEP)", "Postsimulation Activity 2: Home Exercise Program (HEP)", "postsimulation-activity-2-home-exercise-program.docx"),
            ("Presimulation Activity 3: Reflection", "Presimulation Activity 3: Reflection", "presimulation-activity-3-reflection.docx"),
        ),
    },
    {
        "id": "simulation-27", "source": ROOT / "assets" / "Manuscripts" / "E9814_Sim27 Gen Med Checks.docx", "page": ROOT / "pages" / "simulation-27.html", "output_dir": ROOT / "assets" / "downloads" / "simulation-27",
        "activities": (
            ("Presimulation Activity 2: Referral Network", "Presimulation Activity 2: Referral Network", "presimulation-activity-2-referral-network.docx"),
            ("Presimulation Activity 3: Patient History", "Presimulation Activity 3: Patient History", "presimulation-activity-3-patient-history.docx"),
            ("Presimulation Activity 4: Case Presentation", "Presimulation Activity 4: Case Presentation", "presimulation-activity-4-case-presentation.docx"),
            ("Presimulation Activity 5: Plan of Care", "Presimulation Activity 5: Plan of Care", "presimulation-activity-5-plan-of-care.docx"),
            ("Rubric", "Rubric", "rubric.docx"),
            ("Postsimulation Activity 1: After Action Report", "Postsimulation Activity 1: After Action Report", "postsimulation-activity-1-after-action-report.docx"),
            ("Postsimulation Activity 2: Reflection", "Postsimulation Activity 2: Reflection", "postsimulation-activity-2-reflection.docx"),
        ),
    },
    {
        "id": "simulation-28", "source": ROOT / "assets" / "Manuscripts" / "E9814_Sim28 Distracting Rehab.docx", "page": ROOT / "pages" / "simulation-28.html", "output_dir": ROOT / "assets" / "downloads" / "simulation-28",
        "activities": (
            ("Presimulation Activity 2: Therapeutic Intervention for Ankle", "Presimulation Activity 2: Therapeutic Intervention for Ankle", "presimulation-activity-2-therapeutic-intervention-for-ankle.docx"),
            ("Presimulation Activity 3: Warm-Up", "Presimulation Activity 3: Warm-Up", "presimulation-activity-3-warm-up.docx"),
            ("Presimulation Activity 4: Therapeutic Intervention for Elbow", "Presimulation Activity 4: Therapeutic Intervention for Elbow", "presimulation-activity-4-therapeutic-intervention-for-elbow.docx"),
            ("Presimulation Activity 5: Ethics", "Presimulation Activity 5: Ethics", "presimulation-activity-5-ethics.docx"),
            ("Presimulation Activity 6: Policy Development", "Presimulation Activity 6: Policy Development", "presimulation-activity-6-policy-development.docx"),
            ("Presimulation Activity 7: Return to Participation", "Presimulation Activity 7: Return to Participation", "presimulation-activity-7-return-to-participation.docx"),
            ("Rubric", "Rubric", "rubric.docx"),
            ("Postsimulation Activity 1: Reflection", "Postsimulation Activity 1: Reflection", "postsimulation-activity-1-reflection.docx"),
            ("Postsimulation Activity 2: Policies", "Postsimulation Activity 2: Policies", "postsimulation-activity-2-policies.docx"),
        ),
    },
    {
        "id": "simulation-29", "source": ROOT / "assets" / "Manuscripts" / "E9814_Sim29 Evaluation Juggling.docx", "page": ROOT / "pages" / "simulation-29.html", "output_dir": ROOT / "assets" / "downloads" / "simulation-29",
        "activities": (
            ("Presimulation Activity 2: Splinting and Durable Medical Equipment (DME)", "Presimulation Activity 2: Splinting and Durable Medical Equipment (DME)", "presimulation-activity-2-splinting-and-dme.docx"),
            ("Presimulation Activity 3: Concussion Evaluation", "Presimulation Activity 3: Concussion Evaluation", "presimulation-activity-3-concussion-evaluation.docx"),
            ("Presimulation Activity 4: Referral Network", "Presimulation Activity 4: Referral Network", "presimulation-activity-4-referral-network.docx"),
            ("Presimulation Activity 5: Return to Participation", "Presimulation Activity 5: Return to Participation", "presimulation-activity-5-return-to-participation.docx"),
            ("Presimulation Activity 6: Ethics Policy", "Presimulation Activity 6: Ethics Policy", "presimulation-activity-6-ethics-policy.docx"),
            ("Rubric", "Rubric", "rubric.docx"),
            ("Postsimulation Activity 1: Reflection", "Postsimulation Activity 1: Reflection", "postsimulation-activity-1-reflection.docx"),
            ("Postsimulation Activity 2: Documentation", "Postsimulation Activity 2: Documentation", "postsimulation-activity-2-documentation.docx"),
        ),
    },
    {
        "id": "simulation-30", "source": ROOT / "assets" / "Manuscripts" / "E9814_Sim30 Breaking Bad News.docx", "page": ROOT / "pages" / "simulation-30.html", "output_dir": ROOT / "assets" / "downloads" / "simulation-30",
        "activities": (
            ("Presimulation Activity 2: Annotation", "Presimulation Activity 2: Annotation", "presimulation-activity-2-annotation.docx"),
            ("Presimulation Activity 3: Nonverbal Communication", "Presimulation Activity 3: Nonverbal Communication", "presimulation-activity-3-nonverbal-communication.docx"),
            ("Presimulation Activity 4: Active Listening", "Presimulation Activity 4: Active Listening", "presimulation-activity-4-active-listening.docx"),
            ("Presimulation Activity 5: Role-Play", "Presimulation Activity 5: Role-Play", "presimulation-activity-5-role-play.docx"),
            ("Presimulation Activity 6: Research", "Presimulation Activity 6: Research", "presimulation-activity-6-research.docx"),
            ("Postsimulation Activity1: Self-Assessment", "Postsimulation Activity1: Self-Assessment", "postsimulation-activity1-self-assessment.docx"),
            ("Postsimulation Activity 2: Reflection", "Postsimulation Activity 2: Reflection", "postsimulation-activity-2-reflection.docx"),
            ("Postsimulation Activity 3: Documentation", "Postsimulation Activity 3: Documentation", "postsimulation-activity-3-documentation.docx"),
        ),
    },
    {
        "id": "simulation-31", "source": ROOT / "assets" / "Manuscripts" / "E9814_Sim31 Athlete Orientation Day.docx", "page": ROOT / "pages" / "simulation-31.html", "output_dir": ROOT / "assets" / "downloads" / "simulation-31",
        "activities": (
            ("Presimulation Activity 2: Field Research", "Presimulation Activity 2: Field Research", "presimulation-activity-2-field-research.docx"),
            ("Presimulation Activity 3: Critical Appraisal", "Presimulation Activity 3: Critical Appraisal", "presimulation-activity-3-critical-appraisal.docx"),
            ("Presimulation Activity 4: Content Creation", "Presimulation Activity 4: Content Creation", "presimulation-activity-4-content-creation.docx"),
            ("Postsimulation Activity 1: Self-Assessment", "Postsimulation Activity 1: Self-Assessment", "postsimulation-activity-1-self-assessment.docx"),
            ("Postsimulation Activity 2: Reflection", "Postsimulation Activity 2: Reflection", "postsimulation-activity-2-reflection.docx"),
            ("Postsimulation Activity 3: Policy Development", "Postsimulation Activity 3: Policy Development", "postsimulation-activity-3-policy-development.docx"),
        ),
    },
    {
        "id": "simulation-32", "source": ROOT / "assets" / "Manuscripts" / "E9814_Sim32 It all comes together at the hip.docx", "page": ROOT / "pages" / "simulation-32.html", "output_dir": ROOT / "assets" / "downloads" / "simulation-32",
        "activities": (
            ("Presimulation Activity 2: Palpation", "Presimulation Activity 2: Palpation", "presimulation-activity-2-palpation.docx"),
            ("Presimulation Activity 3: Special Test Practice", "Presimulation Activity 3: Special Test Practice", "presimulation-activity-3-special-test-practice.docx"),
            ("Presimulation Activity 4: Modality Selection", "Presimulation Activity 4: Modality Selection", "presimulation-activity-4-modality-selection.docx"),
            ("Presimulation Activity 5: Therapeutic Intervention", "Presimulation Activity 5: Therapeutic Intervention", "presimulation-activity-5-therapeutic-intervention.docx"),
            ("Postsimulation Activity 1: Documentation", "Postsimulation Activity 1: Documentation", "postsimulation-activity-1-documentation.docx"),
            ("Postsimulation Activity 2: Reflection", "Postsimulation Activity 2: Reflection", "postsimulation-activity-2-reflection.docx"),
            ("Postsimulation Activity 3: Home Exercise Program", "Postsimulation Activity 3: Home Exercise Program", "postsimulation-activity-3-home-exercise-program.docx"),
        ),
    },
)


def paragraph_text(element):
    return "".join(element.xpath(".//w:t/text()", namespaces=NS))


def remove_text_prefix(paragraph, prefix):
    remaining = prefix
    for text_node in paragraph.xpath(".//w:t", namespaces=NS):
        value = text_node.text or ""
        if not remaining:
            break
        count = min(len(value), len(remaining))
        if value[:count] != remaining[:count]:
            raise ValueError(f"Expected prefix {prefix!r} in {paragraph_text(paragraph)!r}")
        text_node.text = value[count:]
        remaining = remaining[count:]
    if remaining:
        raise ValueError(f"Could not remove prefix {prefix!r}")


def remove_text_suffix(paragraph, suffix):
    remaining = suffix
    text_nodes = paragraph.xpath(".//w:t", namespaces=NS)
    for text_node in reversed(text_nodes):
        value = text_node.text or ""
        if not remaining:
            break
        count = min(len(value), len(remaining))
        if value[-count:] != remaining[-count:]:
            raise ValueError(f"Expected suffix {suffix!r} in {paragraph_text(paragraph)!r}")
        text_node.text = value[:-count]
        remaining = remaining[:-count]
    if remaining:
        raise ValueError(f"Could not remove suffix {suffix!r}")


def is_begin_marker(text, button_name, title=None, next_text=None):
    if not text.startswith("\\qqBEGIN downloadable content"):
        return False
    marker = text.split("<title>", 1)[0].rstrip("\\")
    if marker == "\\qqBEGIN downloadable content. Button name:":
        return next_text in (f"<title>{title}", f"<b>{title}", f"<a>{title}")
    return marker.endswith(button_name)


def activity_document(source_xml, button_name, title):
    root = etree.fromstring(source_xml)
    body = root.find("w:body", NS)
    children = list(body)
    begin_index = next(
        i for i, node in enumerate(children)
        if is_begin_marker(
            paragraph_text(node),
            button_name,
            title,
            next(
                (
                    paragraph_text(candidate)
                    for candidate in children[i + 1:]
                    if paragraph_text(candidate)
                ),
                None,
            ),
        )
    )
    end_marker = "\\qqEND downloadable content\\"
    end_index = next(
        i for i in range(begin_index + 1, len(children))
        if paragraph_text(children[i]).endswith(end_marker)
        or paragraph_text(children[i]).startswith("\\qqBEGIN downloadable content")
    )
    begin_node = children[begin_index]
    selected = []
    if "<title>" in paragraph_text(begin_node):
        combined_title = deepcopy(begin_node)
        marker_prefix = paragraph_text(begin_node).split("<title>", 1)[0]
        remove_text_prefix(combined_title, marker_prefix)
        selected.append(combined_title)
    for node in children[begin_index + 1:end_index]:
        selected.append(deepcopy(node))
    end_node = children[end_index]
    if paragraph_text(end_node).endswith(end_marker) and paragraph_text(end_node) != end_marker:
        inline_content = deepcopy(end_node)
        remove_text_suffix(inline_content, end_marker)
        selected.append(inline_content)

    title_texts = (f"<title>{title}", f"<b>{title}", f"<a>{title}")
    try:
        title_index = next(i for i, node in enumerate(selected) if paragraph_text(node) in title_texts)
    except StopIteration:
        raise ValueError(f"Download section for {title!r} has no matching title")
    selected = selected[title_index:]
    title_prefix = next(
        prefix for prefix in ("<title>", "<b>", "<a>")
        if paragraph_text(selected[0]).startswith(prefix)
    )
    remove_text_prefix(selected[0], title_prefix)
    for node in selected:
        for prefix in ("<txni>", "<tx>", "<lh>"):
            if paragraph_text(node).startswith(prefix):
                remove_text_prefix(node, prefix)

    section_properties = body.find("w:sectPr", NS)
    for node in list(body):
        body.remove(node)
    for node in selected:
        body.append(node)
    if section_properties is not None:
        body.append(deepcopy(section_properties))
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def write_docx(source_path, output_path, document_xml):
    with ZipFile(source_path) as source, ZipFile(output_path, "w", ZIP_DEFLATED) as output:
        for item in source.infolist():
            payload = document_xml if item.filename == "word/document.xml" else source.read(item.filename)
            output.writestr(item, payload)


def insert_supplied_images(document):
    for paragraph in document.paragraphs:
        marker = paragraph.text.strip()
        if not marker.startswith("\\qqINSERT "):
            continue
        source_name = PureWindowsPath(marker.removeprefix("\\qqINSERT ").strip()).stem
        image_path = IMAGE_DIR / f"{source_name}.png"
        if not image_path.exists():
            raise FileNotFoundError(
                f"No supplied PNG matches image marker {source_name!r}: {image_path}"
            )

        section = document.sections[0]
        available_width = section.page_width - section.left_margin - section.right_margin
        paragraph.clear()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture = paragraph.add_run().add_picture(str(image_path))
        max_width = min(available_width, Inches(6.5))
        max_height = Inches(7.25)
        scale = min(max_width / picture.width, max_height / picture.height, 1)
        picture.width = int(picture.width * scale)
        picture.height = int(picture.height * scale)


def set_run_font(run_element, font_name):
    run_properties = run_element.find(qn("w:rPr"))
    if run_properties is None:
        run_properties = OxmlElement("w:rPr")
        run_element.insert(0, run_properties)
    run_fonts = run_properties.find(qn("w:rFonts"))
    if run_fonts is None:
        run_fonts = OxmlElement("w:rFonts")
        run_properties.insert(0, run_fonts)
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        run_fonts.set(qn(f"w:{attribute}"), font_name)
    for attribute in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "csTheme"):
        run_fonts.attrib.pop(qn(f"w:{attribute}"), None)


def apply_required_font(document):
    normal_style = document.styles["Normal"]
    normal_style.font.name = REQUIRED_FONT
    normal_run_properties = normal_style.element.get_or_add_rPr()
    normal_run_fonts = normal_run_properties.get_or_add_rFonts()
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        normal_run_fonts.set(qn(f"w:{attribute}"), REQUIRED_FONT)

    for part in document.part.package.parts:
        if str(part.partname) == "/word/theme/theme1.xml":
            theme = etree.fromstring(part.blob)
            major_latin = theme.xpath(
                './/*[local-name()="majorFont"]/*[local-name()="latin"]'
            )
            minor_latin = theme.xpath(
                './/*[local-name()="minorFont"]/*[local-name()="latin"]'
            )
            if major_latin:
                major_latin[0].set("typeface", "Aptos Display")
            if minor_latin:
                minor_latin[0].set("typeface", REQUIRED_FONT)
            part._blob = etree.tostring(
                theme, xml_declaration=True, encoding="UTF-8", standalone=True
            )

        element = getattr(part, "element", None)
        if element is None:
            continue
        for run_element in element.iter(qn("w:r")):
            if any(True for _ in run_element.iter(qn("w:t"))):
                set_run_font(run_element, REQUIRED_FONT)


def set_run_size(run_element, half_points):
    run_properties = run_element.find(qn("w:rPr"))
    if run_properties is None:
        run_properties = OxmlElement("w:rPr")
        run_element.insert(0, run_properties)
    for element_name in ("w:sz", "w:szCs"):
        size_element = run_properties.find(qn(element_name))
        if size_element is None:
            size_element = OxmlElement(element_name)
            run_properties.append(size_element)
        size_element.set(qn("w:val"), str(half_points))


def apply_required_heading_size(document):
    heading_style_ids = set()
    for style in document.styles:
        if style.type == WD_STYLE_TYPE.PARAGRAPH and style.name.startswith("Heading"):
            style.font.size = REQUIRED_HEADING_SIZE
            heading_style_ids.add(style.style_id)

    for part in document.part.package.parts:
        element = getattr(part, "element", None)
        if element is None:
            continue
        for paragraph in element.iter(qn("w:p")):
            paragraph_style = paragraph.find(qn("w:pPr"))
            if paragraph_style is None:
                continue
            paragraph_style = paragraph_style.find(qn("w:pStyle"))
            if (
                paragraph_style is None
                or paragraph_style.get(qn("w:val")) not in heading_style_ids
            ):
                continue
            for run_element in paragraph.iter(qn("w:r")):
                if any(True for _ in run_element.iter(qn("w:t"))):
                    set_run_size(run_element, 40)


def finalize_document(output_path):
    document = Document(output_path)
    insert_supplied_images(document)
    title_paragraph = next(paragraph for paragraph in document.paragraphs if paragraph.text)
    try:
        heading_style = document.styles["Heading 1"]
    except KeyError:
        heading_style = document.styles.add_style("Heading 1", WD_STYLE_TYPE.PARAGRAPH)
        heading_style.base_style = document.styles["Normal"]
        heading_style.font.size = Pt(20)
        heading_style.font.color.rgb = RGBColor(0x0F, 0x47, 0x61)
        heading_style.paragraph_format.space_before = Pt(18)
        heading_style.paragraph_format.space_after = Pt(4)
        heading_style.paragraph_format.keep_with_next = True
    title_paragraph.style = heading_style
    for section in document.sections:
        paragraph = section.footer.paragraphs[0]
        paragraph.clear()
        paragraph.style = document.styles["Normal"]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.line_spacing = 1
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        prefix_run = paragraph.add_run(FOOTER_PREFIX)
        title_run = paragraph.add_run(FOOTER_TITLE)
        title_run.italic = True
        suffix_run = paragraph.add_run(FOOTER_SUFFIX)
        for run in (prefix_run, title_run, suffix_run):
            run.font.name = REQUIRED_FONT
            run.font.size = FOOTER_SIZE
    apply_required_font(document)
    apply_required_heading_size(document)
    document.save(output_path)


def main():
    requested = set(sys.argv[1:])
    known = {simulation["id"] for simulation in SIMULATIONS}
    unknown = requested - known
    if unknown:
        raise SystemExit(f"Unknown simulation: {', '.join(sorted(unknown))}")
    selected = [simulation for simulation in SIMULATIONS if not requested or simulation["id"] in requested]
    for simulation in selected:
        source_path = simulation["source"]
        output_dir = simulation["output_dir"]
        output_dir.mkdir(parents=True, exist_ok=True)
        with ZipFile(source_path) as source:
            source_xml = source.read("word/document.xml")
        for button_name, title, filename in simulation["activities"]:
            output_path = output_dir / filename
            document_xml = activity_document(source_xml, button_name, title)
            write_docx(source_path, output_path, document_xml)
            finalize_document(output_path)
            print(f"WROTE: {output_path.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
