"""Check that reader-facing HTML and downloads match Word manuscripts verbatim."""

from html.parser import HTMLParser
from pathlib import Path
import re
import sys
from zipfile import ZipFile
import xml.etree.ElementTree as ET

from docx import Document

from generate_simulation_downloads import FOOTER_TEXT, FOOTER_TITLE, SIMULATIONS, is_begin_marker


ROOT = Path(__file__).resolve().parents[1]
INTRO_SOURCE = ROOT / "assets" / "Manuscripts" / "E9814_FM_introduction.docx"
WORD_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
PRODUCTION_TAG = re.compile(r"^<(?:cn|ct|a|title|txni|tx|lh|b)>")
BEGIN_DOWNLOAD = "\\qqBEGIN downloadable content. Button name: "
END_DOWNLOAD = "\\qqEND downloadable content\\"


def docx_blocks(path):
    with ZipFile(path) as archive:
        document = ET.fromstring(archive.read("word/document.xml"))
    blocks = []
    for paragraph in document.iter(WORD_NS + "p"):
        text = "".join(node.text or "" for node in paragraph.iter(WORD_NS + "t"))
        if text:
            blocks.append(text)
    return blocks


def introduction_blocks():
    return [strip_production_tag(text) for text in docx_blocks(INTRO_SOURCE)]


def strip_production_tag(text):
    if text.strip() == "PHOTO HERE":
        return ""
    if text.startswith("PHOTO HERE") and "<a>" in text:
        return text.split("<a>", 1)[1]
    return PRODUCTION_TAG.sub("", text)


def simulation_page_blocks(source, activities, resources=()):
    blocks = []
    in_download = False
    source_blocks = docx_blocks(source)
    for index, text in enumerate(source_blocks):
        if text.startswith("Navigation menu/button"):
            continue
        if text.startswith("\\qqBEGIN downloadable content"):
            entries = [
                (button, title) for button, title, _filename in activities
            ] + [
                (resource, resource) for resource in resources
            ]
            button_name = next(
                button for button, _title in entries
                if is_begin_marker(
                    text,
                    button,
                    _title,
                    source_blocks[index + 1] if index + 1 < len(source_blocks) else None,
                )
            )
            blocks.append(button_name)
            in_download = True
            continue
        if text == END_DOWNLOAD:
            in_download = False
            continue
        if in_download and text.endswith(END_DOWNLOAD):
            in_download = False
            continue
        if in_download:
            continue
        visible_text = strip_production_tag(text)
        if visible_text.strip():
            blocks.append(visible_text)
    return blocks


def activity_source_blocks(source, button_name, title):
    blocks = []
    collecting = False
    source_blocks = docx_blocks(source)
    for index, text in enumerate(source_blocks):
        if collecting and text.startswith("\\qqBEGIN downloadable content"):
            break
        if is_begin_marker(
            text,
            button_name,
            title,
            source_blocks[index + 1] if index + 1 < len(source_blocks) else None,
        ):
            collecting = True
            if "<title>" in text:
                combined_title = "<title>" + text.split("<title>", 1)[1]
                blocks.append(strip_production_tag(combined_title))
            continue
        if collecting and text.endswith(END_DOWNLOAD):
            inline_content = text.removesuffix(END_DOWNLOAD)
            if inline_content:
                blocks.append(strip_production_tag(inline_content))
            break
        if collecting and not text.startswith(("\\qqINSERT ", "\\qqID:")):
            blocks.append(strip_production_tag(text))
    return blocks


class ManuscriptHTMLParser(HTMLParser):
    block_tags = {"h1", "h2", "h3", "p", "li"}

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.in_manuscript = False
        self.current_tag = None
        self.current_text = []
        self.blocks = []

    def handle_starttag(self, tag, attrs):
        names = {name for name, _ in attrs}
        if "data-manuscript" in names:
            self.in_manuscript = True
        if self.in_manuscript and (tag in self.block_tags or "data-manuscript-block" in names):
            self.current_tag = tag
            self.current_text = []

    def handle_data(self, data):
        if self.current_tag:
            self.current_text.append(data)

    def handle_endtag(self, tag):
        if tag == self.current_tag:
            self.blocks.append("".join(self.current_text))
            self.current_tag = None
            self.current_text = []


def html_blocks(path):
    parser = ManuscriptHTMLParser()
    parser.feed(path.read_text(encoding="utf-8"))
    return parser.blocks


def compare(label, expected, actual):
    if actual == expected:
        print(f"PASS: {label}")
        return True
    print(f"FAIL: {label}")
    count = max(len(expected), len(actual))
    for index in range(count):
        source = expected[index] if index < len(expected) else "<missing>"
        rendered = actual[index] if index < len(actual) else "<missing>"
        if source != rendered:
            print(f"  Block {index + 1}\n    SOURCE: {source!r}\n    OUTPUT: {rendered!r}")
    return False


def validate_footer(path):
    document = Document(path)
    valid = True
    for index, section in enumerate(document.sections, start=1):
        paragraphs = [paragraph for paragraph in section.footer.paragraphs if paragraph.text]
        if len(paragraphs) != 1 or paragraphs[0].text != FOOTER_TEXT:
            print(f"FAIL: {path.relative_to(ROOT).as_posix()} section {index} footer text")
            valid = False
            continue
        if paragraphs[0].alignment != 0:
            print(f"FAIL: {path.relative_to(ROOT).as_posix()} section {index} footer alignment")
            valid = False
        italic_text = "".join(run.text for run in paragraphs[0].runs if run.italic)
        if italic_text != FOOTER_TITLE:
            print(f"FAIL: {path.relative_to(ROOT).as_posix()} section {index} footer italics")
            valid = False
    if valid:
        print(f"PASS: {path.relative_to(ROOT).as_posix()} has the required footer")
    return valid


def validate_heading(path):
    document = Document(path)
    title = next(paragraph for paragraph in document.paragraphs if paragraph.text)
    valid = title.style.name == "Heading 1"
    if valid:
        print(f"PASS: {path.relative_to(ROOT).as_posix()} title uses Heading 1")
    else:
        print(
            f"FAIL: {path.relative_to(ROOT).as_posix()} title style is "
            f"{title.style.name!r}, expected 'Heading 1'"
        )
    return valid


def main():
    requested = set(sys.argv[1:])
    known = {simulation["id"] for simulation in SIMULATIONS}
    unknown = requested - known
    if unknown:
        raise SystemExit(f"Unknown simulation: {', '.join(sorted(unknown))}")
    results = [
        compare(
            f"pages/introduction.html matches {INTRO_SOURCE.name}",
            introduction_blocks(),
            html_blocks(ROOT / "pages" / "introduction.html"),
        ),
    ]
    for simulation in SIMULATIONS:
        if requested and simulation["id"] not in requested:
            continue
        source = simulation["source"]
        page = simulation["page"]
        output_dir = simulation["output_dir"]
        results.append(
            compare(
                f"{page.relative_to(ROOT).as_posix()} matches {source.name}",
                simulation_page_blocks(
                    source,
                    simulation["activities"],
                    simulation.get("resources", ()),
                ),
                html_blocks(page),
            )
        )
        for button_name, _title, filename in simulation["activities"]:
            output_path = output_dir / filename
            results.append(
                compare(
                    f"{output_path.relative_to(ROOT).as_posix()} matches its manuscript section",
                    activity_source_blocks(source, button_name, _title),
                    docx_blocks(output_path),
                )
            )
            results.append(validate_footer(output_path))
            results.append(validate_heading(output_path))
    return 0 if all(results) else 1


if __name__ == "__main__":
    sys.exit(main())
